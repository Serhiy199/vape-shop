# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt


OUTPUT_PATH = r"D:\Copy_WSL_Project\vape-shop\Vape_Shop_Report_2026-04-18_fixed.docx"


def configure_styles(document: Document) -> None:
    styles = document.styles

    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Normal"].font.size = Pt(12)

    styles["Title"].font.name = "Times New Roman"
    styles["Title"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    styles["Title"].font.size = Pt(18)
    styles["Title"].font.bold = True

    for style_name, size in (("Heading 1", 14), ("Heading 2", 13)):
        styles[style_name].font.name = "Times New Roman"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.bold = True


def add_heading_1(document: Document, text: str) -> None:
    document.add_paragraph(text, style="Heading 1")


def add_heading_2(document: Document, text: str) -> None:
    document.add_paragraph(text, style="Heading 2")


def add_bullet(document: Document, text: str) -> None:
    document.add_paragraph(f"• {text}")


doc = Document()
configure_styles(doc)

title = doc.add_paragraph("Доповідь про виконану роботу за 18.04.2026", style="Title")
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph("Проєкт: Vape Shop Platform")

add_heading_1(doc, "Мета роботи за день")
doc.add_paragraph(
    "Завершити Phase 3: Auth And Access Control для проєкту Vape Shop Platform, "
    "реалізувати повний credentials-based auth flow, роль ADMIN, seed admin-акаунта "
    "та перевірити end-to-end сценарії доступу для клієнта й адміністратора."
)

add_heading_1(doc, "Виконані роботи")

add_heading_2(doc, "1. Підготовка та стабілізація auth skeleton")
add_bullet(doc, "Доведено до робочого стану існуючий auth skeleton на Auth.js.")
add_bullet(doc, "Переведено session strategy на JWT для стабільної роботи з route protection.")
add_bullet(doc, "Узгоджено Auth.js config, session callbacks, JWT callbacks і типізацію next-auth.")
add_bullet(doc, "Оновлено захист маршрутів під Next.js 16 через proxy замість застарілого middleware convention.")

add_heading_2(doc, "2. Реєстрація клієнта")
add_bullet(doc, "Доопрацьовано POST /api/auth/register з валідацією через Zod.")
add_bullet(doc, "Реалізовано хешування пароля через bcryptjs.")
add_bullet(doc, "Зафіксовано, що реєстрація завжди створює користувача з роллю CLIENT.")
add_bullet(doc, "Створено UI-сторінку /register і клієнтську форму реєстрації з success/error flow.")

add_heading_2(doc, "3. Вхід, logout і сесія")
add_bullet(doc, "Реалізовано login через Credentials Provider (email + password).")
add_bullet(doc, "Створено окрему сторінку /login і форму входу через server action.")
add_bullet(doc, "Реалізовано logout через signOut() з поверненням на головну сторінку.")
add_bullet(doc, "Додано відображення активної сесії в UI: користувач, email, роль, статус авторизації.")

add_heading_2(doc, "4. Захист приватних маршрутів")
add_bullet(doc, "Захищено /account/* для авторизованих користувачів.")
add_bullet(doc, "Захищено /admin/* виключно для користувачів з роллю ADMIN.")
add_bullet(doc, "Реалізовано захист у два шари: proxy на периметрі та server-side guards у layout-ах.")
add_bullet(doc, "Додано контрольні сторінки /account, /admin і /admin/users для перевірки доступів.")

add_heading_2(doc, "5. Роль ADMIN як окремий сценарій")
add_bullet(doc, "Централізовано рольову логіку в окремих helper-ах.")
add_bullet(doc, "Налаштовано role-aware redirects після login: ADMIN переходить у /admin, CLIENT — на storefront.")
add_bullet(doc, "Оновлено UI так, щоб ADMIN і CLIENT бачили різні дії та різні домашні точки входу.")
add_bullet(doc, "Роль у сесії використовується не лише для guard-ів, а і для відображення статусу користувача в інтерфейсі.")

add_heading_2(doc, "6. Seed admin-акаунта та фінальна перевірка")
add_bullet(doc, "Підтверджено роботу prisma seed для створення або оновлення admin-акаунта.")
add_bullet(doc, "Оновлено README з інструкціями для prisma:seed та auth smoke test.")
add_bullet(doc, "Проведено end-to-end перевірку через локальний dev server.")
add_bullet(doc, "Підтверджено сценарії доступу для guest, CLIENT і ADMIN.")

add_heading_1(doc, "Результат на кінець дня")
add_bullet(doc, "Клієнт може зареєструватися і увійти в систему.")
add_bullet(doc, "Logout працює коректно.")
add_bullet(doc, "Сесія працює і відображається в UI.")
add_bullet(doc, "/account/* доступний тільки авторизованим користувачам.")
add_bullet(doc, "/admin/* доступний тільки користувачам з роллю ADMIN.")
add_bullet(doc, "Seed admin-акаунта працює та може використовуватися для стартового доступу до адмінки.")
add_bullet(doc, "Phase 3 завершено повністю згідно з поставленою метою.")

add_heading_1(doc, "Що ще не реалізовано")
add_bullet(doc, "Повний каталог і бізнес-логіка storefront.")
add_bullet(doc, "CRUD-модулі адмінки для категорій, товарів, замовлень і promo codes.")
add_bullet(doc, "Wishlist, cart, checkout і order flow.")
add_bullet(doc, "Інтеграція оплат та подальша e-commerce логіка.")
add_bullet(doc, "Розширені приватні кабінети клієнта з реальними даними.")

add_heading_1(doc, "Висновок")
doc.add_paragraph(
    "У межах Phase 3 побудовано повний фундамент авторизації та контролю доступу "
    "для Vape Shop Platform. Система вже підтримує реєстрацію клієнта, логін, logout, "
    "активну сесію, role-based access control, окремий сценарій для ADMIN та стартову "
    "перевірену адмінську зону. Проєкт готовий до переходу до наступної фази, де можна "
    "розвивати бізнес-модулі storefront і адміністративну панель поверх уже стабільної "
    "auth-основи."
)

doc.save(OUTPUT_PATH)
print(OUTPUT_PATH)
